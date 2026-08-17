import io
import os
import json
import re
from typing import Dict, Any, List
from PyPDF2 import PdfReader
import docx
import requests
from dotenv import load_dotenv

import schemas

load_dotenv()

# ==========================================
# 1. DOCUMENT TEXT EXTRACTION
# ==========================================

def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs if para.text])
        return text
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

def extract_text(file_bytes: bytes, filename: str) -> str:
    if filename.lower().endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    elif filename.lower().endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    return ""

# ==========================================
# 2. ROBUST TEXT CHUNKER (Built-in + LangChain)
# ==========================================

def chunk_text(text: str, chunk_size: int = 500, chunk_overlap: int = 80) -> List[str]:
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        return [doc.page_content for doc in splitter.create_documents([text])]
    except Exception:
        pass

    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        return [doc.page_content for doc in splitter.create_documents([text])]
    except Exception:
        pass

    # Native Python semantic chunker
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    current_chunk = ""
    for p in paragraphs:
        if len(current_chunk) + len(p) < chunk_size:
            current_chunk += "\n" + p
        else:
            if current_chunk:
                chunks.append(current_chunk.strip())
            current_chunk = p
    if current_chunk:
        chunks.append(current_chunk.strip())
    return chunks if chunks else [text]

# ==========================================
# 3. MULTI-MODEL GEMINI INVOCATION ENGINE
# ==========================================

ACTIVE_GEMINI_MODELS = [
    "gemini-flash-lite-latest",
    "gemini-flash-latest",
    "gemini-2.5-flash-lite",
    "gemma-4-26b-a4b-it"
]

def _call_gemini_with_fallback(api_key: str, prompt: str) -> str:
    """Tries active Gemini models with automatic fallback if one is busy (503/429)."""
    last_err = None
    for model_name in ACTIVE_GEMINI_MODELS:
        try:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"
            payload = {
                "contents": [{"parts": [{"text": prompt}]}],
                "generationConfig": {"temperature": 0.2, "maxOutputTokens": 2048}
            }
            resp = requests.post(url, json=payload, timeout=25)
            if resp.status_code == 200:
                data = resp.json()
                cand = data.get("candidates", [{}])[0]
                text = cand.get("content", {}).get("parts", [{}])[0].get("text", "")
                if text.strip():
                    print(f"[AI Service] Successfully generated insights using {model_name}")
                    return text
            else:
                print(f"[AI Service] Model {model_name} returned status {resp.status_code}")
                last_err = f"Status {resp.status_code}: {resp.text[:120]}"
        except Exception as e:
            print(f"[AI Service] Model {model_name} error: {e}")
            last_err = str(e)
            
    raise RuntimeError(f"All Gemini models exhausted. Last error: {last_err}")

# ==========================================
# 4. JSON RESPONSE PARSER & CLEANER
# ==========================================

def _clean_and_parse_json(raw_text: str) -> Dict[str, Any]:
    cleaned = raw_text.strip()
    if "```json" in cleaned:
        cleaned = cleaned.split("```json")[1].split("```")[0].strip()
    elif "```" in cleaned:
        cleaned = cleaned.split("```")[1].split("```")[0].strip()

    first_brace = cleaned.find("{")
    last_brace = cleaned.rfind("}")
    if first_brace != -1 and last_brace != -1:
        cleaned = cleaned[first_brace:last_brace + 1]

    return json.loads(cleaned)

# ==========================================
# 5. FALLBACK HEURISTIC GENERATOR
# ==========================================

def generate_fallback_heuristic_analysis(text: str) -> Dict[str, Any]:
    if not text.strip():
        return {
            "overall_score": 0,
            "ats_score": 0,
            "skill_match": 0,
            "issues_found": 5,
            "ai_summary": "Empty resume file uploaded. No text could be extracted.",
            "ats_feedback": "Document contains no readable text layers. Ensure standard PDF/DOCX export.",
            "action_verb_feedback": "No action verbs detected.",
            "bullet_suggestions": [],
            "missing_keywords": ["Experience", "Skills", "Education", "Projects"],
            "strengths": [],
            "improvements": ["Upload a non-empty, text-searchable resume file."]
        }
        
    text_lower = text.lower()
    words = text_lower.split()
    word_count = len(words)
    
    sections = ['experience', 'education', 'skills', 'projects', 'summary', 'objective']
    found_sections = [sec for sec in sections if sec in text_lower]
    section_score = min(100, (len(found_sections) / 4) * 100)
    
    action_verbs = ['developed', 'managed', 'created', 'led', 'designed', 'implemented', 'improved', 'increased', 'reduced', 'optimized', 'spearheaded', 'orchestrated']
    found_verbs = [verb for verb in action_verbs if verb in text_lower]
    verb_score = min(100, (len(found_verbs) / 5) * 100)
    
    has_numbers = bool(re.search(r'\d+', text))
    has_percentages = bool(re.search(r'\d+%', text))
    has_currency = bool(re.search(r'\$\d+', text))
    
    metrics_score = 0
    if has_numbers: metrics_score += 40
    if has_percentages: metrics_score += 30
    if has_currency: metrics_score += 30
    
    length_score = 100
    if word_count < 200: length_score = 50
    elif word_count > 1000: length_score = 70
        
    overall_score = round(min(max((section_score * 0.3) + (verb_score * 0.3) + (metrics_score * 0.2) + (length_score * 0.2), 20), 96), 1)
    ats_score = round(min(max((section_score * 0.6) + (length_score * 0.4), 25), 98), 1)
    skill_match = round(min(max((verb_score * 0.4) + (section_score * 0.6), 20), 95), 1)
    
    issues_count = 0
    if len(found_sections) < 4: issues_count += 1
    if len(found_verbs) < 4: issues_count += 1
    if not has_numbers: issues_count += 1
    if word_count < 250 or word_count > 900: issues_count += 1
    if not (has_percentages or has_currency): issues_count += 1

    return {
        "overall_score": overall_score,
        "ats_score": ats_score,
        "skill_match": skill_match,
        "issues_found": max(issues_count, 1),
        "ai_summary": f"Resume analysis extracted {word_count} words across key sections ({', '.join(found_sections) if found_sections else 'general'}). The document showcases relevant background but would benefit from greater metric quantification and targeted keyword expansion.",
        "ats_feedback": f"Standard section headers detected: {len(found_sections)}/6. Ensure clean single-column hierarchy for optimal ATS parsing.",
        "action_verb_feedback": f"Found {len(found_verbs)} high-impact action verbs. Aim to start every bullet with strong velocity verbs like 'Architected', 'Spearheaded', or 'Optimized'.",
        "bullet_suggestions": [
            {
                "original": "Responsible for managing tasks and working with the engineering team.",
                "improved": "Spearheaded Agile sprint execution with 6 engineers, accelerating feature shipping cadence by 30%.",
                "reason": "Replaces passive duty statement with quantified business impact."
            }
        ],
        "missing_keywords": ["CI/CD Pipelines", "System Architecture", "KPI Tracking", "Cross-Functional Leadership"],
        "strengths": [
            f"Clear section division with {len(found_sections)} standard resume sections.",
            "Adequate overall word count volume for scanning.",
            "Logical chronological presentation."
        ],
        "improvements": [
            "Add quantifiable business metrics (%, $, volume) to every project bullet.",
            "Incorporate more industry-specific technical keywords.",
            "Replace passive phrasing with active leadership verbs."
        ]
    }

# ==========================================
# 6. MAIN RAG + GEMINI LLM ANALYSIS PIPELINE
# ==========================================

def analyze_resume_with_rag(text: str) -> Dict[str, Any]:
    api_key = os.getenv("GOOGLE_API_KEY", "").strip()
    if not api_key:
        print("[AI Service] No GOOGLE_API_KEY set in backend/.env. Using fallback analysis.")
        return generate_fallback_heuristic_analysis(text)

    # 1. Chunk document
    chunks = chunk_text(text, chunk_size=500, chunk_overlap=80)

    # 2. Vector Retrieval (RAG)
    rag_context = ""
    try:
        from langchain_google_genai import GoogleGenerativeAIEmbeddings
        from langchain_community.vectorstores import FAISS
        from langchain_core.documents import Document

        doc_objs = [Document(page_content=c) for c in chunks]
        embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001", google_api_key=api_key)
        vectorstore = FAISS.from_documents(doc_objs, embeddings)
        retriever = vectorstore.as_retriever(search_kwargs={"k": 4})

        q1 = retriever.invoke("Work experience accomplishments metrics percentages leadership")
        q2 = retriever.invoke("Skills technical tools technologies certifications education")
        q3 = retriever.invoke("Key projects bullet points achievements responsibilities")

        all_retrieved = {d.page_content for d in (q1 + q2 + q3)}
        rag_context = "\n---\n".join(all_retrieved)
        print("[AI Service] RAG vector context retrieved successfully via FAISS.")
    except Exception as rag_err:
        print(f"[AI Service] RAG vector retrieval note: {rag_err}. Using context chunks.")
        rag_context = "\n---\n".join(chunks[:6])

    # 3. Formulate Prompt
    analysis_prompt = f"""You are an elite Executive Resume Reviewer, ATS Algorithm Specialist, and Technical Hiring Manager.
Analyze the candidate's resume thoroughly using the extracted context below.
CRITICAL INSTRUCTION: Your feedback MUST be highly personalized. Do NOT provide generic advice like "Use the Google XYZ formula" or "Quantify your impact" or "Fix character spacing". You MUST reference specific projects, specific skills, or specific sentences from the candidate's actual resume text. If you suggest an improvement, explain exactly which sentence or section in their resume needs it.

Resume Sections Context:
\"\"\"
{rag_context}
\"\"\"

Full Resume Text Sample:
\"\"\"
{text[:3500]}
\"\"\"

Analyze the resume and return a STRICT, VALID JSON object with the exact keys below. Do NOT output any preamble or extra text.

JSON Schema:
{{
  "overall_score": <number between 0 and 100>,
  "ats_score": <number between 0 and 100>,
  "skill_match": <number between 0 and 100>,
  "issues_found": <integer count of issues>,
  "ai_summary": "<Detailed executive summary evaluating profile strength, career readiness, and competitive edge>",
  "ats_feedback": "<Specific feedback on layout, formatting, headers, and parsing readability>",
  "action_verb_feedback": "<Critique on action verb velocity and proactive achievement phrasing>",
  "bullet_suggestions": [
    {{
      "original": "<A real weak or passive line found in the resume>",
      "improved": "<High-impact rewritten bullet with strong action verb and quantified metric/results>",
      "reason": "<Why this change is stronger>"
    }}
  ],
  "missing_keywords": ["<Specific Keyword 1 missing from their stack>", "<Specific Keyword 2>"],
  "strengths": ["<Specific Strength 1 found in their text>", "<Specific Strength 2>"],
  "improvements": [
    "<Hyper-specific Improvement 1 referencing their actual project/role>", 
    "<Hyper-specific Improvement 2 referencing their actual skills/text>"
  ],
  "parsed_ats_data": [
    {{
      "field": "First Name",
      "status": "found",
      "value": "<Extracted First Name or 'Missing'>"
    }},
    {{
      "field": "Last Name",
      "status": "found",
      "value": "<Extracted Last Name or 'Missing'>"
    }},
    {{
      "field": "Email",
      "status": "found",
      "value": "<Extracted Email or 'Missing'>"
    }},
    {{
      "field": "Phone",
      "status": "found",
      "value": "<Extracted Phone or 'Missing'>"
    }},
    {{
      "field": "LinkedIn URL",
      "status": "found",
      "value": "<Extracted LinkedIn URL or 'Missing'>"
    }},
    {{
      "field": "Education (Degree)",
      "status": "found",
      "value": "<Extracted Highest Degree or 'Missing'>"
    }},
    {{
      "field": "Graduation Year",
      "status": "found",
      "value": "<Extracted Year or 'Missing'>"
    }},
    {{
      "field": "Work Experience 1",
      "status": "found",
      "value": "<Extracted Latest Job Title and Company or 'Missing'>"
    }}
  ]
}}
"""

    # 4. Invoke LLM with fallback
    try:
        print("[AI Service] Invoking Gemini LLM for resume analysis...")
        raw_response = _call_gemini_with_fallback(api_key, analysis_prompt)
        parsed_data = _clean_and_parse_json(raw_response)

        # Validate with Pydantic
        validated = schemas.LLMAnalysisOutput(**parsed_data)
        print("[AI Service] Gemini LLM analysis completed and validated successfully!")
        return validated.model_dump()

    except Exception as llm_err:
        print(f"[AI Service] LLM generation error: {llm_err}. Falling back to heuristic engine.")
        return generate_fallback_heuristic_analysis(text)

# ==========================================
# 7. INTERVIEW PREP GENERATOR
# ==========================================

def generate_interview_questions(analysis: Any) -> Dict[str, Any]:
    api_key = os.getenv("GOOGLE_API_KEY", "").strip()
    if not api_key:
        raise ValueError("No GOOGLE_API_KEY set for interview generation.")

    # Build context from analysis
    context_str = f"AI Summary: {analysis.ai_summary}\n"
    if analysis.strengths:
        context_str += f"Strengths: {analysis.strengths}\n"
    if analysis.parsed_ats_data:
        context_str += f"Extracted Data: {analysis.parsed_ats_data}\n"
    if analysis.missing_keywords:
        context_str += f"Missing/Suggested Skills: {analysis.missing_keywords}\n"
    if analysis.improvements:
        context_str += f"Improvements: {analysis.improvements}\n"
    if analysis.bullet_suggestions:
        context_str += f"Bullet points context: {analysis.bullet_suggestions}\n"

    prompt = f"""You are an elite Technical Hiring Manager conducting an interview based on the candidate's resume.
Your task is to generate EXACTLY 3 highly personalized interview questions for this specific candidate.

Candidate Resume Context:
\"\"\"
{context_str}
\"\"\"

Follow these strict rules:
1. Question 1 (type: "resume_project"): Based directly on something mentioned in the resume (e.g., a project, work experience, achievement, or technical implementation).
2. Question 2 (type: "technical"): Based on one or more technical skills detected from the resume context. Do NOT invent unrelated skills.
3. Question 3 (type: "job_role"): Infer the target job role from the resume context and ask a role-specific question that matches the candidate's skills and the inferred role's standard requirements.

Output the result strictly as a valid JSON object matching this schema:
{{
  "questions": [
    {{
      "type": "resume_project",
      "question": "...",
      "reason": "..."
    }},
    {{
      "type": "technical",
      "question": "...",
      "reason": "..."
    }},
    {{
      "type": "job_role",
      "question": "...",
      "reason": "..."
    }}
  ]
}}

Do NOT output any additional text, preamble, or markdown blocks, only the JSON. 
Make the questions reasonable for a real technical interview (Medium to Medium/Hard difficulty). 
Ensure the questions are distinct and highly specific to the candidate.
"""
    try:
        raw_response = _call_gemini_with_fallback(api_key, prompt)
        parsed_data = _clean_and_parse_json(raw_response)
        
        # Validate structure
        if "questions" not in parsed_data or len(parsed_data["questions"]) != 3:
            # Fallback if structure is violated
            print("[AI Service] Invalid questions length, attempting repair.")
            questions = parsed_data.get("questions", [])
            while len(questions) < 3:
                questions.append({
                    "type": "technical",
                    "question": "Based on your resume, how do you handle complex technical problem-solving?",
                    "reason": "General technical problem solving assessment."
                })
            parsed_data["questions"] = questions[:3]

        validated = schemas.InterviewPrepResponse(**parsed_data)
        return validated.model_dump()
    except Exception as e:
        print(f"[AI Service] Interview Generation Error: {e}")
        raise RuntimeError("Failed to generate interview questions.")

# ==========================================
# 8. DETERMINISTIC RESUME vs JD MATCHING ENGINE
# ==========================================

SKILL_ONTOLOGY = {
    "AWS": ["aws", "amazon web services", "s3", "ec2", "lambda", "sagemaker", "rds", "dynamodb", "cloudwatch", "iam", "ecs", "eks", "fargate", "sqs", "sns", "redshift", "cloudfront"],
    "CI/CD": ["ci/cd", "cicd", "continuous integration", "continuous deployment", "continuous delivery", "github actions", "gitlab ci", "jenkins", "circleci", "argo cd"],
    "TypeScript": ["typescript", "type script", "ts"],
    "JavaScript": ["javascript", "js", "ecmascript", "node", "nodejs", "express"],
    "Python": ["python", "py", "fastapi", "flask", "django", "numpy", "pandas"],
    "Docker": ["docker", "containerization", "containers", "dockerfile", "docker-compose"],
    "Kubernetes": ["kubernetes", "k8s", "kubectl", "helm"],
    "GraphQL": ["graphql", "graph ql", "apollo graphql"],
    "REST API": ["rest", "restful", "rest api", "api development", "web services", "json api"],
    "SQL": ["sql", "postgresql", "postgres", "mysql", "sqlite", "tsql", "plsql", "database design"],
    "NoSQL": ["nosql", "mongodb", "mongo", "dynamodb", "redis", "cassandra"],
    "React": ["react", "react.js", "reactjs", "next.js", "nextjs", "redux"],
    "Vue": ["vue", "vue.js", "vuejs", "nuxt"],
    "Angular": ["angular", "angularjs"],
    "Java": ["java", "spring", "spring boot"],
    "C++": ["c++", "cpp"],
    "C#": ["c#", ".net", "dotnet", "asp.net"],
    "Go": ["golang", "go programming"],
    "Rust": ["rust"],
    "Agile": ["agile", "agile methodology", "scrum", "kanban", "sprint planning"],
    "Git": ["git", "github", "gitlab", "version control", "bitbucket"],
    "Machine Learning": ["machine learning", "ml", "deep learning", "tensorflow", "pytorch", "scikit-learn", "sklearn", "ai", "artificial intelligence"],
    "MLOps": ["mlops", "mlflow", "kubeflow", "wandb", "model monitoring", "model deployment"],
    "Microservices": ["microservices", "distributed systems", "service-oriented architecture"],
    "System Design": ["system design", "system architecture", "scalable systems"],
    "GCP": ["gcp", "google cloud", "google cloud platform", "bigquery"],
    "Azure": ["azure", "microsoft azure"],
    "Linux": ["linux", "bash", "shell scripting", "ubuntu", "unix"]
}

def extract_resume_sections(text: str) -> Dict[str, List[str]]:
    """Parses resume text into logical section lines."""
    lines = text.split("\n")
    sections = {
        "technical_skills": [],
        "experience": [],
        "projects": [],
        "summary": [],
        "education": [],
        "certifications": [],
        "other": []
    }
    
    current_sec = "other"
    
    sec_patterns = {
        "technical_skills": re.compile(r'^(technical\s+skills|skills|technologies|tech\text|tools|competencies|programming)', re.I),
        "experience": re.compile(r'^(work\s+experience|professional\s+experience|experience|employment\s+history|history)', re.I),
        "projects": re.compile(r'^(projects|key\s+projects|academic\s+projects|personal\s+projects)', re.I),
        "summary": re.compile(r'^(summary|professional\s+summary|profile|about\s+me|objective)', re.I),
        "education": re.compile(r'^(education|academic\s+background|degrees)', re.I),
        "certifications": re.compile(r'^(certifications|licenses|courses)', re.I)
    }

    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line:
            continue
            
        matched_header = False
        if len(cleaned_line) < 45 and not cleaned_line.endswith("."):
            for sec_name, pattern in sec_patterns.items():
                if pattern.search(cleaned_line):
                    current_sec = sec_name
                    matched_header = True
                    break
        
        if not matched_header:
            sections[current_sec].append(cleaned_line)
            
    return sections

def extract_jd_required_skills(jd_text: str) -> List[str]:
    """Extracts required tech keywords from Job Description using Ontology & NLP pattern matching."""
    extracted = set()
    jd_lower = jd_text.lower()
    
    # 1. Match against known skill ontology & aliases
    for canonical_name, aliases in SKILL_ONTOLOGY.items():
        for alias in aliases:
            # Word boundary regex search
            escaped_alias = re.escape(alias)
            pattern = rf'\b{escaped_alias}\b' if len(alias) <= 4 else rf'\b{escaped_alias}'
            if re.search(pattern, jd_lower, re.IGNORECASE):
                extracted.add(canonical_name)
                break

    # 2. Extract uppercase/capitalized tech words from JD (e.g., GraphQL, Terraform, Kafka)
    potential_words = re.findall(r'\b[A-Z][A-Za-z0-9+#.-]{2,15}\b', jd_text)
    ignore_words = {"The", "And", "With", "For", "You", "Our", "We", "Will", "Role", "Job", "Team", "Work", "Status", "Years", "Senior", "Junior", "Lead", "Full", "Stack", "Engineer", "Developer", "Manager", "Company"}
    for word in potential_words:
        if word not in ignore_words and len(word) > 2:
            # Normalize canonical name if matches ontology
            matched = False
            for canonical, aliases in SKILL_ONTOLOGY.items():
                if word.lower() in [a.lower() for a in aliases] or word.lower() == canonical.lower():
                    extracted.add(canonical)
                    matched = True
                    break
            if not matched and word.isupper() or word in {"GraphQL", "Terraform", "Kafka", "Postgres", "Redis", "Next.js", "Tailwind"}:
                extracted.add(word)

    return sorted(list(extracted)) if extracted else ["AWS", "CI/CD", "TypeScript", "Python", "Agile"]

def perform_deterministic_skill_match(resume_text: str, jd_text: str) -> Dict[str, Any]:
    """
    Deterministic Skill Matcher & 4-Tier Classifier:
    - STRONG_MATCH: Skill/alias present in Experience or Projects with active evidence
    - MATCH: Skill/alias explicitly listed in Technical Skills or Experience
    - WEAK_MATCH: Skill/alias present only in Summary or Education
    - MISSING: Skill not found in any resume section
    """
    sections = extract_resume_sections(resume_text)
    required_skills = extract_jd_required_skills(jd_text)
    
    section_display_names = {
        "technical_skills": "Technical Skills",
        "experience": "Experience",
        "projects": "Projects",
        "summary": "Summary",
        "education": "Education",
        "certifications": "Certifications",
        "other": "General Content"
    }

    strong_matches = []
    matches = []
    weak_matches = []
    missing_keywords = []
    
    score_sum = 0.0

    for skill in required_skills:
        aliases = SKILL_ONTOLOGY.get(skill, [skill.lower()])
        if skill.lower() not in [a.lower() for a in aliases]:
            aliases.append(skill.lower())

        locations = []
        evidence = []
        found_in_exp_or_proj = False
        found_in_tech_skills = False
        found_in_other = False

        for sec_key, lines in sections.items():
            display_name = section_display_names[sec_key]
            for line in lines:
                line_lower = line.lower()
                for alias in aliases:
                    escaped_alias = re.escape(alias)
                    pattern = rf'\b{escaped_alias}\b' if len(alias) <= 4 else rf'\b{escaped_alias}'
                    if re.search(pattern, line_lower, re.IGNORECASE):
                        if display_name not in locations:
                            locations.append(display_name)
                        if line not in evidence and len(evidence) < 2:
                            evidence.append(line[:160])
                        
                        if sec_key in ("experience", "projects"):
                            found_in_exp_or_proj = True
                        elif sec_key == "technical_skills":
                            found_in_tech_skills = True
                        else:
                            found_in_other = True
                        break

        # Classify tier
        if found_in_exp_or_proj and (len(evidence) > 0 or found_in_tech_skills):
            status = "strong_match"
            confidence = 0.98
            item = {
                "skill": skill,
                "status": status,
                "locations": locations,
                "evidence": evidence if evidence else [f"Demonstrated across {', '.join(locations)}."],
                "confidence": confidence
            }
            strong_matches.append(item)
            score_sum += 100.0

        elif found_in_exp_or_proj or found_in_tech_skills:
            status = "match"
            confidence = 0.90
            item = {
                "skill": skill,
                "status": status,
                "locations": locations,
                "evidence": evidence if evidence else [f"Listed in {', '.join(locations)}."],
                "confidence": confidence
            }
            matches.append(item)
            score_sum += 85.0

        elif found_in_other:
            status = "weak_match"
            confidence = 0.70
            item = {
                "skill": skill,
                "status": status,
                "locations": locations,
                "evidence": evidence if evidence else [f"Mentioned in {', '.join(locations)}."],
                "confidence": confidence
            }
            weak_matches.append(item)
            score_sum += 50.0

        else:
            missing_keywords.append(skill)
            score_sum += 0.0

    total_skills = len(required_skills)
    match_score = round(score_sum / total_skills, 1) if total_skills > 0 else 75.0

    # Deterministic Recommendations (Rule: NEVER recommend adding detected skills!)
    recommendations = []
    for item in weak_matches:
        recommendations.append(
            f"Strengthen '{item['skill']}': It is mentioned in {', '.join(item['locations'])}, but adding a quantified achievement bullet in your Work Experience will raise it to a Strong Match."
        )

    for skill_name in missing_keywords[:3]:
        recommendations.append(
            f"Consider adding '{skill_name}': If you have hands-on experience with {skill_name}, incorporate it into your Technical Skills and relevant project bullets."
        )

    if not recommendations:
        recommendations.append("Excellent alignment! Your resume strongly evidences all core technical requirements for this job role.")

    return {
        "match_score": match_score,
        "strong_matches": strong_matches,
        "matches": matches,
        "weak_matches": weak_matches,
        "missing_keywords": missing_keywords,
        "recommendations": recommendations
    }

def analyze_job_match(resume_text: str, jd_text: str, job_title: str = "") -> Dict[str, Any]:
    """Combines deterministic skill matcher as ground truth with Gemini LLM for executive summary."""
    # 1. Run Deterministic Matcher (Source of Truth)
    det_results = perform_deterministic_skill_match(resume_text, jd_text)
    
    # 2. Invoke Gemini LLM strictly for explanation synthesis
    api_key = os.getenv("GOOGLE_API_KEY", "").strip()
    summary = ""
    
    if api_key:
        prompt = f"""You are an Executive Recruiter and ATS Specialist.
Below are the GROUND-TRUTH deterministic skill match results comparing candidate resume to Job Description for '{job_title or "Target Role"}':

Overall Match Score: {det_results['match_score']}%
Strong Matches: {[m['skill'] for m in det_results['strong_matches']]}
Matches: {[m['skill'] for m in det_results['matches']]}
Weak Matches: {[m['skill'] for m in det_results['weak_matches']]}
Missing Keywords: {det_results['missing_keywords']}

Instruction: Write a concise 2-3 sentence executive summary explaining the candidate's alignment. 
CRITICAL RULE: Respect the ground truth above. Do NOT state that AWS, CI/CD, or TypeScript are missing if they are listed under Strong Matches or Matches.

Return ONLY a plain text paragraph with no JSON or markdown tags.
"""
        try:
            summary = _call_gemini_with_fallback(api_key, prompt).strip()
        except Exception:
            summary = ""

    if not summary:
        strong_count = len(det_results['strong_matches']) + len(det_results['matches'])
        missing_count = len(det_results['missing_keywords'])
        summary = f"Your profile exhibits strong technical alignment with {strong_count} key job requirements matched. Addressing the {missing_count} missing keywords will optimize your resume for ATS screening."

    return {
        "match_score": det_results["match_score"],
        "summary": summary,
        "strong_matches": det_results["strong_matches"],
        "matches": det_results["matches"],
        "weak_matches": det_results["weak_matches"],
        "missing_keywords": det_results["missing_keywords"],
        "recommendations": det_results["recommendations"]
    }

